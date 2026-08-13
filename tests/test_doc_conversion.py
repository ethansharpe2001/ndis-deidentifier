import os
import subprocess
import sys
import tempfile

import pytest
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app import converters  # noqa: E402
from app.converters import DocConversionError, convert_doc_to_docx, find_soffice  # noqa: E402
from app.deidentify import deidentify_file  # noqa: E402
from tests.make_sample import build_sample  # noqa: E402
from tests.pii_fixtures import EXPECTED_TAGS, FORBIDDEN_STRINGS  # noqa: E402

_SOFFICE = find_soffice()


def test_convert_raises_clear_error_when_libreoffice_missing(monkeypatch):
    monkeypatch.setattr(converters, "find_soffice", lambda: None)
    with tempfile.TemporaryDirectory() as tmp:
        fake_doc = os.path.join(tmp, "plan.doc")
        with open(fake_doc, "wb") as f:
            f.write(b"not a real .doc, doesn't matter for this test")
        with pytest.raises(DocConversionError, match=r"LibreOffice"):
            convert_doc_to_docx(fake_doc, os.path.join(tmp, "out"))


def test_deidentify_file_on_doc_raises_clear_error_when_libreoffice_missing(monkeypatch):
    monkeypatch.setattr(converters, "find_soffice", lambda: None)
    with tempfile.TemporaryDirectory() as tmp:
        fake_doc = os.path.join(tmp, "plan.doc")
        with open(fake_doc, "wb") as f:
            f.write(b"not a real .doc, doesn't matter for this test")
        with pytest.raises(DocConversionError, match=r"LibreOffice"):
            deidentify_file(fake_doc)


@pytest.mark.skipif(_SOFFICE is None, reason="LibreOffice not installed on this machine")
def test_real_doc_round_trip_removes_all_planted_pii():
    """End-to-end test using a real LibreOffice install: build a .docx
    fixture, convert it down to legacy .doc with LibreOffice (the reverse of
    what the app does, just to get a realistic .doc fixture), then feed that
    .doc through the app and confirm the planted PII is gone.
    """
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "sample.docx")
        build_sample(docx_path)

        proc = subprocess.run(
            [_SOFFICE, "--headless", "--norestore", "--convert-to", "doc", "--outdir", tmp, docx_path],
            capture_output=True,
            text=True,
            timeout=120,
        )
        doc_path = os.path.join(tmp, "sample.doc")
        assert proc.returncode == 0 and os.path.exists(doc_path), (
            f"Failed to produce .doc fixture via LibreOffice: {proc.stderr or proc.stdout}"
        )

        output_path = os.path.join(tmp, "sample_deidentified.docx")
        result = deidentify_file(doc_path, output_path)
        assert result.total > 0

        redacted = Document(output_path)
        text = "\n".join(
            [p.text for p in redacted.paragraphs]
            + [cell.text for t in redacted.tables for row in t.rows for cell in row.cells]
        )
        for forbidden in FORBIDDEN_STRINGS:
            assert forbidden not in text, f"Leaked PII survived redaction: {forbidden!r}"
        for tag in EXPECTED_TAGS:
            assert tag in text, f"Expected placeholder missing: {tag!r}"
