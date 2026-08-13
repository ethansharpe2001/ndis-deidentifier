"""Builds synthetic NDIS Behaviour Support Plan fixtures (.docx and .txt),
structured the way real templates/exports typically are, so the
de-identification pipeline can be exercised end-to-end without any real
participant data. Both builders plant the same identifying strings so tests
can share one FORBIDDEN_STRINGS / EXPECTED_TAGS list across formats.
"""
from docx import Document


def build_sample(path: str) -> None:
    doc = Document()

    doc.add_heading("NDIS Behaviour Support Plan", level=0)

    doc.add_heading("Participant Details", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Participant Name", "Sarah Jennifer Mitchell"),
        ("Date of Birth", "14/03/2009"),
        ("NDIS Number", "430128754"),
        ("Address", "27 Wattle Grove, Bendigo VIC 3550"),
        ("Medicare Number", "2950 12092 1"),  # valid mod-10 Medicare checksum
        ("Contact Phone", "0412 345 678"),
        ("Contact Email", "sarah.mitchell@example.com"),
    ]
    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    doc.add_heading("Guardian / Family Contact", level=1)
    table2 = doc.add_table(rows=0, cols=2)
    table2.style = "Table Grid"
    for label, value in [
        ("Parent/Guardian Name", "David Mitchell"),
        ("Emergency Contact Phone", "0400 111 222"),
    ]:
        row_cells = table2.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    doc.add_heading("Plan Details", level=1)
    table3 = doc.add_table(rows=0, cols=2)
    table3.style = "Table Grid"
    for label, value in [
        ("Behaviour Support Practitioner", "Dr. Amanda Chen"),
        ("Provider Organisation", "Bright Pathways Allied Health"),
        ("Plan Date", "1 July 2026"),
        ("Plan Review Date", "1 January 2027"),
    ]:
        row_cells = table3.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    doc.add_heading("Background", level=1)
    doc.add_paragraph(
        "Sarah Jennifer Mitchell is a 17-year-old participant who lives with her "
        "father, David Mitchell, at their home in Bendigo. Sarah was referred to "
        "Bright Pathways Allied Health following an increase in the frequency of "
        "behaviours of concern at school."
    )
    doc.add_paragraph(
        "Dr. Amanda Chen conducted a functional behaviour assessment with Sarah "
        "over three sessions. Sarah's teacher reported that Sarah often becomes "
        "distressed during transitions between classes."
    )

    doc.add_heading("Strategies", level=1)
    doc.add_paragraph(
        "Staff supporting Sarah should provide a five-minute warning before any "
        "transition. If Sarah becomes distressed, staff should contact David "
        "Mitchell on 0400 111 222 or email the practitioner at "
        "amanda.chen@brightpathways.example.com."
    )

    doc.add_heading("Sign-off", level=1)
    doc.add_paragraph("Practitioner Signature: Dr. Amanda Chen")
    doc.add_paragraph("Date: 1 July 2026")

    doc.save(path)


def build_sample_txt(path: str) -> None:
    """Plain-text equivalent of `build_sample`, using the "Label: Value" and
    split "Label:\\nValue" layouts real .txt exports use instead of tables.
    """
    lines = [
        "NDIS Behaviour Support Plan",
        "",
        "== Participant Details ==",
        "Participant Name:",
        "Sarah Jennifer Mitchell",
        "Date of Birth: 14/03/2009",
        "NDIS Number: 430128754",
        "Address: 27 Wattle Grove, Bendigo VIC 3550",
        "Medicare Number: 2950 12092 1",
        "Contact Phone: 0412 345 678",
        "Contact Email: sarah.mitchell@example.com",
        "",
        "== Guardian / Family Contact ==",
        "Parent/Guardian Name: David Mitchell",
        "Emergency Contact Phone: 0400 111 222",
        "",
        "== Plan Details ==",
        "Behaviour Support Practitioner: Dr. Amanda Chen",
        "Provider Organisation: Bright Pathways Allied Health",
        "Plan Date: 1 July 2026",
        "Plan Review Date: 1 January 2027",
        "",
        "== Background ==",
        "Sarah Jennifer Mitchell is a 17-year-old participant who lives with her "
        "father, David Mitchell, at their home in Bendigo. Sarah was referred to "
        "Bright Pathways Allied Health following an increase in the frequency of "
        "behaviours of concern at school.",
        "Dr. Amanda Chen conducted a functional behaviour assessment with Sarah "
        "over three sessions. Sarah's teacher reported that Sarah often becomes "
        "distressed during transitions between classes.",
        "",
        "== Strategies ==",
        "Staff supporting Sarah should provide a five-minute warning before any "
        "transition. If Sarah becomes distressed, staff should contact David "
        "Mitchell on 0400 111 222 or email the practitioner at "
        "amanda.chen@brightpathways.example.com.",
        "",
        "== Sign-off ==",
        "Practitioner Signature: Dr. Amanda Chen",
        "Date: 1 July 2026",
        "",
    ]
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


if __name__ == "__main__":
    import sys

    build_sample(sys.argv[1] if len(sys.argv) > 1 else "sample_bsp.docx")
    print("Sample written.")
