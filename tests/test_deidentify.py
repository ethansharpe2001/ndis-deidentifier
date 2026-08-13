import os
import sys
import tempfile

import pytest
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.deidentify import deidentify_file, deidentify_files, default_output_path  # noqa: E402
from tests.make_sample import build_sample, build_sample_txt  # noqa: E402
from tests.pii_fixtures import EXPECTED_TAGS, FORBIDDEN_STRINGS  # noqa: E402


def _full_text(docx_path: str) -> str:
    doc = Document(docx_path)
    chunks = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def test_deidentify_removes_all_planted_pii():
    with tempfile.TemporaryDirectory() as tmp:
        sample_path = os.path.join(tmp, "sample.docx")
        output_path = os.path.join(tmp, "sample_deidentified.docx")
        build_sample(sample_path)

        result = deidentify_file(sample_path, output_path)
        assert result.total > 0

        text = _full_text(output_path)

        for forbidden in FORBIDDEN_STRINGS:
            assert forbidden not in text, f"Leaked PII survived redaction: {forbidden!r}"

        for tag in EXPECTED_TAGS:
            assert tag in text, f"Expected placeholder missing: {tag!r}"

        # Structure must survive: same number of tables/rows, headings intact.
        original = Document(sample_path)
        redacted = Document(output_path)
        assert len(original.tables) == len(redacted.tables)
        for ot, rt in zip(original.tables, redacted.tables):
            assert len(ot.rows) == len(rt.rows)
        assert len(original.paragraphs) == len(redacted.paragraphs)


def test_unsupported_extension_rejected():
    with tempfile.TemporaryDirectory() as tmp:
        bad_path = os.path.join(tmp, "not_supported.pdf")
        with open(bad_path, "w") as f:
            f.write("hello")
        with pytest.raises(ValueError, match=r"Unsupported file type"):
            deidentify_file(bad_path)


def test_extensionless_file_rejected():
    with tempfile.TemporaryDirectory() as tmp:
        bad_path = os.path.join(tmp, "no_extension")
        with open(bad_path, "w") as f:
            f.write("hello")
        with pytest.raises(ValueError, match=r"Unsupported file type"):
            deidentify_file(bad_path)


def test_corrupt_docx_raises_instead_of_silently_producing_garbage():
    with tempfile.TemporaryDirectory() as tmp:
        bad_path = os.path.join(tmp, "corrupt.docx")
        with open(bad_path, "wb") as f:
            f.write(b"this is not a real docx file, just garbage bytes")
        with pytest.raises(Exception):
            deidentify_file(bad_path)


def test_default_output_path_adds_suffix_and_keeps_extension():
    assert default_output_path("plan.docx") == "plan_deidentified.docx"
    assert default_output_path("plan.txt") == "plan_deidentified.txt"


def test_default_output_path_doc_input_becomes_docx_output():
    # Legacy .doc can only be converted to .docx on the way in - there's no
    # way to write back to the original binary format.
    assert default_output_path("plan.doc") == "plan_deidentified.docx"


def test_batch_one_bad_file_does_not_stop_the_rest():
    with tempfile.TemporaryDirectory() as tmp:
        good_docx = os.path.join(tmp, "good.docx")
        good_txt = os.path.join(tmp, "good.txt")
        bad_path = os.path.join(tmp, "bad.pdf")
        missing_path = os.path.join(tmp, "does_not_exist.docx")
        build_sample(good_docx)
        build_sample_txt(good_txt)
        with open(bad_path, "w") as f:
            f.write("hello")

        outcomes = deidentify_files([good_docx, good_txt, bad_path, missing_path])
        by_input = {o.input_path: o for o in outcomes}

        assert by_input[good_docx].error is None
        assert by_input[good_docx].result.total > 0
        assert os.path.exists(by_input[good_docx].output_path)

        assert by_input[good_txt].error is None
        assert by_input[good_txt].result.total > 0
        assert os.path.exists(by_input[good_txt].output_path)

        assert by_input[bad_path].error is not None
        assert by_input[bad_path].output_path is None

        assert by_input[missing_path].error is not None
        assert by_input[missing_path].output_path is None
